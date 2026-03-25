"""Tests for src/document_editor.py."""
from __future__ import annotations

import re
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from src.document_editor import (
    DocumentEditError,
    apply_tailoring_plan,
    backup_resume,
    extract_resume_text,
    _replace_paragraph_runs,
    _replace_paragraph_text,
)
from src.models import TailoringPlan


# ---------------------------------------------------------------------------
# backup_resume
# ---------------------------------------------------------------------------


def test_backup_creates_timestamped_file(tmp_resume_path: str) -> None:
    backup_path = backup_resume(tmp_resume_path)
    assert Path(backup_path).exists()
    # Filename must contain _backup_ followed by a timestamp
    assert re.search(r"_backup_\d{8}_\d{6}\.docx$", backup_path)


def test_backup_raises_for_missing_file() -> None:
    with pytest.raises(FileNotFoundError):
        backup_resume("/nonexistent/path/resume.docx")


# ---------------------------------------------------------------------------
# extract_resume_text
# ---------------------------------------------------------------------------


def test_extract_resume_text_returns_all_paragraphs(tmp_resume_path: str) -> None:
    text = extract_resume_text(tmp_resume_path)
    assert "Python" in text
    assert "SQL" in text
    assert "PyTorch" in text


def test_extract_resume_text_raises_for_missing_file() -> None:
    with pytest.raises(FileNotFoundError):
        extract_resume_text("/nonexistent/resume.docx")


# ---------------------------------------------------------------------------
# _replace_paragraph_text
# ---------------------------------------------------------------------------


def test_replace_paragraph_text_matches_and_replaces(tmp_resume_path: str) -> None:
    doc = Document(tmp_resume_path)
    old = "Optimized SQL queries reducing report latency by 40%"
    new = "Optimized SQL queries and Python ETL jobs, cutting latency by 55%"
    result = _replace_paragraph_text(doc, old, new)
    assert result is True
    texts = [p.text for p in doc.paragraphs]
    assert new in texts
    assert old not in texts


def test_replace_paragraph_text_no_match_returns_false(tmp_resume_path: str) -> None:
    doc = Document(tmp_resume_path)
    result = _replace_paragraph_text(doc, "this text does not exist", "replacement")
    assert result is False


# ---------------------------------------------------------------------------
# _replace_paragraph_runs — formatting preservation
# ---------------------------------------------------------------------------


def test_replace_preserves_bold_formatting(tmp_path: object) -> None:
    """Replacing a bold run must produce a bold run."""
    path = tmp_path / "bold_test.docx"  # type: ignore[operator]
    doc = Document()
    para = doc.add_paragraph("Original bold bullet")
    para.runs[0].bold = True
    para.runs[0].font.size = Pt(12)
    doc.save(str(path))

    doc2 = Document(str(path))
    target_para = doc2.paragraphs[0]
    _replace_paragraph_runs(target_para, "New bold bullet")

    # The replacement run must still be bold
    assert target_para.runs[0].bold is True
    assert target_para.text == "New bold bullet"


def test_replace_preserves_font_size(tmp_path: object) -> None:
    path = tmp_path / "font_test.docx"  # type: ignore[operator]
    doc = Document()
    para = doc.add_paragraph("Original text")
    para.runs[0].font.size = Pt(11)
    doc.save(str(path))

    doc2 = Document(str(path))
    _replace_paragraph_runs(doc2.paragraphs[0], "Replaced text")
    assert doc2.paragraphs[0].runs[0].font.size == Pt(11)


# ---------------------------------------------------------------------------
# apply_tailoring_plan
# ---------------------------------------------------------------------------


def test_apply_tailoring_plan_produces_valid_docx(
    tmp_resume_path: str, tmp_path: object
) -> None:
    output = str(tmp_path / "tailored.docx")  # type: ignore[operator]
    plan = TailoringPlan(
        bullet_replacements=[
            (
                "Optimized SQL queries reducing report latency by 40%",
                "Optimized PostgreSQL queries and Python ETL jobs, reducing latency by 55%",
            )
        ],
        skills_to_emphasize=["PyTorch", "SQL"],
    )
    result_path = apply_tailoring_plan(
        tmp_resume_path, plan, output, backup=False
    )
    assert Path(result_path).exists()
    doc = Document(result_path)
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "PostgreSQL" in full_text


def test_apply_tailoring_plan_replaces_summary(
    tmp_resume_path: str, tmp_path: object
) -> None:
    output = str(tmp_path / "tailored_summary.docx")  # type: ignore[operator]
    new_summary = (
        "ML-focused CS student with PyTorch, SQL, and Python expertise targeting "
        "Summer 2026 ML Engineering internships."
    )
    plan = TailoringPlan(summary_replacement=new_summary)
    apply_tailoring_plan(tmp_resume_path, plan, output, backup=False)
    doc = Document(output)
    assert doc.paragraphs[0].text == new_summary


def test_apply_tailoring_plan_raises_for_missing_resume(tmp_path: object) -> None:
    output = str(tmp_path / "out.docx")  # type: ignore[operator]
    with pytest.raises(FileNotFoundError):
        apply_tailoring_plan("/no/such/file.docx", TailoringPlan(), output)


def test_apply_tailoring_plan_creates_backup(
    tmp_resume_path: str, tmp_path: object
) -> None:
    output = str(tmp_path / "tailored.docx")  # type: ignore[operator]
    resume_dir = Path(tmp_resume_path).parent
    apply_tailoring_plan(tmp_resume_path, TailoringPlan(), output, backup=True)
    backups = list(resume_dir.glob("*_backup_*.docx"))
    assert len(backups) == 1
