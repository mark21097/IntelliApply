"""Shared pytest fixtures for IntelliApply tests."""
from __future__ import annotations

import pytest
from docx import Document
from docx.shared import Pt

from src.models import JobDescription


@pytest.fixture()
def tmp_resume_path(tmp_path: object) -> str:
    """Create a minimal .docx resume for Alex Kim and return its path.

    The resume contains a summary paragraph and several bullet points that
    can be targeted by document_editor tests.
    """
    path = tmp_path / "test_resume.docx"  # type: ignore[operator]
    doc = Document()

    # Summary paragraph
    summary = doc.add_paragraph(
        "Motivated CS student with strong Python, SQL, and ML skills "
        "seeking Summer 2026 internship."
    )
    summary.runs[0].bold = False

    # Section heading (not a bullet — used to test non-replacement)
    doc.add_paragraph("EXPERIENCE")

    # Bullet points (bold run to test formatting preservation)
    bullet1 = doc.add_paragraph(
        "Built data ingestion pipeline in Python processing 500K rows/day"
    )
    bullet1.runs[0].bold = True
    bullet1.runs[0].font.size = Pt(11)

    doc.add_paragraph(
        "Optimized SQL queries reducing report latency by 40%"
    )

    doc.add_paragraph(
        "Developed ML classification model using PyTorch achieving 94% accuracy"
    )

    doc.add_paragraph("SKILLS")
    doc.add_paragraph("Python, SQL, PyTorch, scikit-learn, NumPy, pandas, Git")

    doc.save(str(path))
    return str(path)


@pytest.fixture()
def mock_job_ml_intern() -> JobDescription:
    """JobDescription fixture for a Summer 2026 ML Engineering internship."""
    return JobDescription(
        raw_text=(
            "NeuralCore AI is hiring a Machine Learning Engineering Intern for "
            "Summer 2026. You will build and evaluate ML pipelines using Python, "
            "PyTorch, and scikit-learn. Strong understanding of algorithms, data "
            "structures, and SQL required. Experience with NLP or computer vision "
            "is a plus. We value candidates who can write clean, well-tested code."
        ),
        title="Machine Learning Engineering Intern",
        company="NeuralCore AI",
        required_skills=["Python", "PyTorch", "scikit-learn", "SQL", "algorithms"],
        url="https://neuralcore.ai/jobs/ml-intern-summer-2026",
    )


@pytest.fixture()
def mock_job_swe_intern() -> JobDescription:
    """JobDescription fixture for a Summer 2026 Software Engineering internship."""
    return JobDescription(
        raw_text=(
            "DataStack Inc. is looking for a Software Engineering Intern (Summer 2026) "
            "to join our platform team. You will design and implement REST APIs in "
            "Python, write SQL queries against our PostgreSQL data warehouse, and "
            "contribute to our data pipeline infrastructure. Proficiency in Python, "
            "Git, and SQL is required. Familiarity with cloud platforms (AWS/GCP) is "
            "a plus."
        ),
        title="Software Engineering Intern",
        company="DataStack Inc.",
        required_skills=["Python", "SQL", "REST APIs", "Git", "PostgreSQL"],
        url="https://datastack.io/careers/swe-intern-2026",
    )
