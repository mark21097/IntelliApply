"""python-docx XML manipulation for resume tailoring.

Handles all .docx read/write operations, preserving paragraph-level formatting
(indentation, spacing, numbering) while replacing run-level text content.
"""
from __future__ import annotations

import copy
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from src.models import TailoringPlan


class DocumentEditError(Exception):
    """Raised when a .docx editing operation fails."""


def backup_resume(resume_path: str) -> str:
    """Copy the original resume to a timestamped backup in the same directory.

    Args:
        resume_path: Absolute or relative path to the source .docx file.

    Returns:
        The path to the newly created backup file.

    Raises:
        FileNotFoundError: If resume_path does not exist.
    """
    src = Path(resume_path)
    if not src.exists():
        raise FileNotFoundError(f"Resume not found: {resume_path}")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = src.parent / f"{src.stem}_backup_{timestamp}{src.suffix}"
    shutil.copy2(src, backup_path)
    return str(backup_path)


def extract_resume_text(resume_path: str) -> str:
    """Extract all paragraph text from a .docx file.

    Args:
        resume_path: Path to the .docx file.

    Returns:
        Newline-separated plain text of all paragraphs.

    Raises:
        FileNotFoundError: If resume_path does not exist.
        DocumentEditError: If the file cannot be parsed as a .docx.
    """
    if not Path(resume_path).exists():
        raise FileNotFoundError(f"Resume not found: {resume_path}")
    try:
        doc = Document(resume_path)
    except Exception as exc:
        raise DocumentEditError(f"Failed to open .docx: {resume_path}") from exc
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())


def apply_tailoring_plan(
    resume_path: str,
    plan: TailoringPlan,
    output_path: str,
    backup: bool = True,
) -> str:
    """Apply a TailoringPlan to a resume and save the result.

    Applies bullet replacements first, then the summary replacement (if any).
    Paragraph-level formatting (indent, spacing, numbering) is preserved.
    Run-level formatting (bold, italic, font) is copied from the first run
    of each replaced paragraph.

    Args:
        resume_path: Path to the source .docx file.
        plan: LLM-generated tailoring instructions.
        output_path: Destination path for the tailored .docx.
        backup: If True, back up the original before editing.

    Returns:
        The output_path where the tailored resume was saved.

    Raises:
        FileNotFoundError: If resume_path does not exist.
        DocumentEditError: If the document cannot be parsed or saved.
    """
    if not Path(resume_path).exists():
        raise FileNotFoundError(f"Resume not found: {resume_path}")

    if backup:
        backup_resume(resume_path)

    try:
        doc = Document(resume_path)
    except Exception as exc:
        raise DocumentEditError(f"Failed to open .docx: {resume_path}") from exc

    # Apply bullet-level replacements
    for old_text, new_text in plan.bullet_replacements:
        _replace_paragraph_text(doc, old_text, new_text)

    # Apply summary replacement to the first non-empty paragraph
    if plan.summary_replacement:
        _replace_summary(doc, plan.summary_replacement)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(output_path)
    except Exception as exc:
        raise DocumentEditError(f"Failed to save .docx: {output_path}") from exc

    return output_path


def _replace_paragraph_text(doc: Document, old_text: str, new_text: str) -> bool:
    """Find the first paragraph containing old_text and replace its content.

    Paragraph-level formatting (pPr) is preserved. Run formatting is copied
    from the first run of the matched paragraph.

    Args:
        doc: Open python-docx Document object.
        old_text: Substring to search for across all paragraph runs.
        new_text: Replacement text.

    Returns:
        True if a replacement was made, False if old_text was not found.
    """
    for para in doc.paragraphs:
        if old_text in para.text:
            _replace_paragraph_runs(para, new_text)
            return True
    return False


def _replace_summary(doc: Document, new_text: str) -> None:
    """Replace the first non-empty paragraph (typically the summary/objective).

    Args:
        doc: Open python-docx Document object.
        new_text: Replacement text for the summary paragraph.
    """
    for para in doc.paragraphs:
        if para.text.strip():
            _replace_paragraph_runs(para, new_text)
            return


def _replace_paragraph_runs(paragraph: object, new_text: str) -> None:
    """Replace all runs in a paragraph with a single run containing new_text.

    Preserves the paragraph element (pPr: indentation, spacing, numbering).
    Copies run-level formatting (rPr: bold, italic, font name, font size)
    from the first existing run, if one is present.

    Args:
        paragraph: A python-docx Paragraph object.
        new_text: The replacement text content.
    """
    p = paragraph._p  # type: ignore[attr-defined]

    # Snapshot the paragraph properties and first run's formatting before clearing
    pPr = p.find(qn("w:pPr"))
    pPr_copy = copy.deepcopy(pPr) if pPr is not None else None

    existing_runs = p.findall(qn("w:r"))
    rPr_copy = None
    if existing_runs:
        rPr = existing_runs[0].find(qn("w:rPr"))
        rPr_copy = copy.deepcopy(rPr) if rPr is not None else None

    # Clear all child elements from the paragraph XML element
    p.clear()

    # Restore paragraph properties
    if pPr_copy is not None:
        p.append(pPr_copy)

    # Build a new <w:r> run element with restored formatting and new text
    from lxml import etree

    r = etree.SubElement(p, qn("w:r"))
    if rPr_copy is not None:
        r.append(rPr_copy)
    t = etree.SubElement(r, qn("w:t"))
    t.text = new_text
    # Preserve leading/trailing spaces
    if new_text != new_text.strip():
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
